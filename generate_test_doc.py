#!/usr/bin/env python3
"""
Generate a Word test document for Reshma.

Usage:
    python generate_test_doc.py              # creates fix_verification_test.docx
    python generate_test_doc.py myfile.docx  # custom output name
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Test data ──────────────────────────────────────────────────────────────────

HEADER = {
    "title":       "Garden Swap — Fix Verification Test Cases",
    "tester":      "Reshma Rajkumar",
    "date":        "_______________",
    "environment": "https://garden-swap.onrender.com",
    "logins": [
        "priya@demo.com / demo1234",
        "faizan@demo.com / demo1234",
        "(or sign up your own account)",
    ],
    "note": (
        "These 8 cases verify the bugs found in Stage 2 testing. "
        "Please fill in Pass/Fail and add notes for anything unexpected."
    ),
}

COLUMNS = ["ID", "Feature", "Steps", "Expected Result", "Pass/Fail", "Notes"]

TEST_CASES = [
    {
        "id": "TC-F01",
        "feature": "Demo credentials work",
        "steps": (
            "1. Open the site in a fresh private/incognito window.\n"
            "2. Click Sign In → enter priya@demo.com / demo1234.\n"
            "3. Click Sign In."
        ),
        "expected": (
            "Logged in successfully. Sign Out button appears, + FAB visible, "
            "feed shows 18 plant listings."
        ),
    },
    {
        "id": "TC-F02",
        "feature": "Login modal closes cleanly",
        "steps": (
            "1. Sign out (if signed in).\n"
            "2. Click Sign In, enter valid credentials, click Sign In."
        ),
        "expected": (
            "Modal closes completely. Feed is visible. "
            "No modal overlay remains on screen."
        ),
    },
    {
        "id": "TC-F03",
        "feature": "Stale session auto-logout",
        "steps": (
            "1. Sign in as any user.\n"
            "2. Open browser DevTools (F12) → Application tab → "
            "Local Storage → set gs_token to the text: Bearer invalidtoken123\n"
            "3. Click the + FAB to create a new listing."
        ),
        "expected": (
            "Automatically signed out. Sign In modal opens. "
            "No 'User not found' error shown."
        ),
    },
    {
        "id": "TC-F04",
        "feature": "Create a listing",
        "steps": (
            "1. Sign in with any account.\n"
            "2. Click + FAB.\n"
            "3. Fill in title, plant type, condition, upload any photo.\n"
            "4. Click Submit."
        ),
        "expected": (
            "Modal closes. Your new listing appears in the feed."
        ),
    },
    {
        "id": "TC-F05",
        "feature": "Submit a giveaway swap request",
        "steps": (
            "1. Sign in as faizan@demo.com / demo1234.\n"
            "2. Open any listing NOT owned by Faizan.\n"
            "3. Click Request Swap.\n"
            "4. Select 'Request as giveaway', add a message, click Submit."
        ),
        "expected": (
            "Success alert appears. Redirected to Chats tab. "
            "Swap conversation is listed."
        ),
    },
    {
        "id": "TC-F06",
        "feature": "Chats tab loads",
        "steps": (
            "1. After TC-F05, click the Chats tab."
        ),
        "expected": (
            "Conversation card visible with listing title, "
            "other party name, and 'Pending' state badge."
        ),
    },
    {
        "id": "TC-F07",
        "feature": "Open chat and send a message",
        "steps": (
            "1. Click the conversation from TC-F06.\n"
            "2. Type a message in the input box.\n"
            "3. Press Enter or click Send."
        ),
        "expected": (
            "Message appears as a right-aligned bubble in the chat."
        ),
    },
    {
        "id": "TC-F08",
        "feature": "Accept swap and confirm handoff",
        "steps": (
            "1. Open a second browser/incognito window.\n"
            "2. Sign in as the lister of the listing from TC-F05 "
            "(the demo user who owns that plant).\n"
            "3. Go to Chats → open the pending swap → click Accept.\n"
            "4. Both parties click 'Confirm Handoff Complete' "
            "(switch windows to do both sides)."
        ),
        "expected": (
            "State changes to Completed. "
            "Rate button appears for both parties."
        ),
    },
]

SUMMARY_ROWS = [
    ["Total Test Cases", "8"],
    ["Passed", ""],
    ["Failed", ""],
    ["Blocked", ""],
]


# ── Styling helpers ─────────────────────────────────────────────────────────────

GREEN  = RGBColor(0x2D, 0x6A, 0x4F)
LIGHT  = RGBColor(0xD8, 0xF3, 0xDC)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREY   = RGBColor(0xF8, 0xF9, 0xFA)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def bold_cell(cell, text, size=10, color=WHITE):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color


def normal_cell(cell, text, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)


# ── Document builder ────────────────────────────────────────────────────────────

def build_doc(out_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # Title
    title = doc.add_heading(HEADER["title"], level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = GREEN

    # Meta block
    doc.add_paragraph(f"Tester: {HEADER['tester']}")
    doc.add_paragraph(f"Date: {HEADER['date']}")
    doc.add_paragraph(f"Environment: {HEADER['environment']}")
    logins_para = doc.add_paragraph("Demo logins: ")
    logins_para.add_run(", ".join(HEADER["logins"])).italic = True
    doc.add_paragraph(HEADER["note"])
    doc.add_paragraph("")

    # Test cases table
    col_widths = [Inches(0.6), Inches(1.1), Inches(2.2), Inches(1.9), Inches(0.7), Inches(1.4)]
    tbl = doc.add_table(rows=1, cols=len(COLUMNS))
    tbl.style = "Table Grid"

    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, col in enumerate(COLUMNS):
        bold_cell(hdr_cells[i], col, size=9)
        set_cell_bg(hdr_cells[i], GREEN)

    # Data rows
    for idx, tc in enumerate(TEST_CASES):
        row_cells = tbl.add_row().cells
        bg = GREY if idx % 2 == 0 else WHITE

        normal_cell(row_cells[0], tc["id"], size=9)
        normal_cell(row_cells[1], tc["feature"], size=9)
        normal_cell(row_cells[2], tc["steps"], size=9)
        normal_cell(row_cells[3], tc["expected"], size=9)
        row_cells[4].text = ""  # Pass/Fail — tester fills in
        row_cells[5].text = ""  # Notes — tester fills in

        for i, cell in enumerate(row_cells):
            set_cell_bg(cell, bg)
            cell.width = col_widths[i]

    doc.add_paragraph("")

    # Summary heading
    doc.add_heading("Test Summary", level=2)

    # Summary table
    stbl = doc.add_table(rows=len(SUMMARY_ROWS), cols=2)
    stbl.style = "Table Grid"
    for i, (label, val) in enumerate(SUMMARY_ROWS):
        cells = stbl.rows[i].cells
        bold_cell(cells[0], label, size=10, color=GREEN if i == 0 else RGBColor(0,0,0))
        normal_cell(cells[1], val, size=10)
        set_cell_bg(cells[0], LIGHT if i > 0 else GREEN)
        if i == 0:
            stbl.rows[0].cells[0]._tc.get_or_add_tcPr()
            bold_cell(cells[0], label, size=10, color=WHITE)
            set_cell_bg(cells[0], GREEN)

    doc.add_paragraph("")
    sig = doc.add_paragraph("Tester Signature: _______________     Date: _______________")
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(out_path)
    print(f"Written: {out_path}")


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).parent / "fix_verification_test.docx"
    build_doc(out)
